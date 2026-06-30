import io
from docx import Document


def llenar_visita(template_path: str, datos, telefono: str) -> io.BytesIO:
    """Carga la plantilla de visita técnica y rellena la tabla con los datos del asegurado.

    Args:
        template_path: ruta al .docx plantilla.
        datos: pd.Series con las columnas del Excel (Num_Siniestro, Nro_Carpeta, etc.).
        telefono: teléfono del asegurado ingresado por el inspector.

    Returns:
        BytesIO con el documento generado.
    """
    doc = Document(template_path)
    tabla = doc.tables[0]

    mapping = {
        "Número siniestro":            str(datos["Num_Siniestro"]),
        "Numero carpeta":              str(datos["Nro_Carpeta"]),
        "Dirección Riesgo asegurado":  str(datos["Dirección Riesgo Asegurado"]),
        "Asegurado":                   str(datos["Asegurado"]),
        "RUT":                         str(datos["Rut"]),
        "Teléfono":                    telefono,
    }

    for row in tabla.rows:
        label = row.cells[0].text.strip()
        if label in mapping:
            cell = row.cells[1]
            p = cell.paragraphs[0]
            valor = mapping[label]
            if p.runs:
                p.runs[0].text = valor
                for r in p.runs[1:]:
                    r.text = ""
            else:
                p.add_run(valor)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
