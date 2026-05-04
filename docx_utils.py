import io
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT

LOGO_PATH = "logo.png"


def _keep_con_siguiente(parrafo):
    """Evita que el párrafo quede solo al final de una página (Keep with next)."""
    pPr = parrafo._p.get_or_add_pPr()
    kn = OxmlElement("w:keepNext")
    pPr.append(kn)


def _titulo_seccion(doc, texto: str):
    """Agrega un título de sección con espacio previo y 'keep with next'."""
    p = doc.add_paragraph()
    _keep_con_siguiente(p)
    p.paragraph_format.space_before = Pt(16)
    r = p.add_run(texto)
    r.bold = True
    r.underline = True
    return p


def generar_documento(
    datos,
    detalle_visita: str,
    hay_inspeccion: bool,
    inspecciones: list,
    sin_danos: bool,
    danos: list,
    observaciones: list,
    imagenes_data: list,
    imagen_error: str = "",
) -> io.BytesIO:
    """Genera el informe Word y lo devuelve como BytesIO."""
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(14)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    # Pie de página
    p = doc.sections[0].footer.paragraphs[0]
    run = p.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Borde de página
    sectPr = doc.sections[0]._sectPr
    pgBorders = OxmlElement("w:pgBorders")
    pgBorders.set(qn("w:offsetFrom"), "page")
    for side in ["top", "left", "bottom", "right"]:
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "6")
        b.set(qn("w:space"), "24")
        b.set(qn("w:color"), "000000")
        pgBorders.append(b)
    sectPr.append(pgBorders)

    # ─── Página 1: Datos (página propia) ────────────────────────────────────
    p_logo = doc.add_paragraph()
    p_logo.add_run().add_picture(LOGO_PATH, width=Inches(2))
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    titulo = doc.add_paragraph()
    r = titulo.add_run("INFORME DE SINIESTRO")
    r.bold = True
    r.font.size = Pt(16)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")

    subt = doc.add_paragraph()
    r = subt.add_run("1. Datos del Siniestro:")
    r.bold = True
    r.underline = True

    tabla = doc.add_table(rows=0, cols=2)
    tabla.style = "Table Grid"
    for label, key in [
        ("Numero de Carpeta",           "Nro_Carpeta"),
        ("Numero de Siniestro",         "Num_Siniestro"),
        ("Dirección Riesgo Asegurado",  "Dirección Riesgo Asegurado"),
        ("Nombre Asegurado",            "Asegurado"),
        ("Rut Asegurado",               "Rut"),
    ]:
        row = tabla.add_row().cells
        row[0].text = label
        row[1].text = str(datos[key])

    # Salto de página solo después de la sección 1
    doc.add_page_break()

    # Contador dinámico de secciones (sección 1 = Datos, siempre)
    sec = 1

    # ─── Sección 2: Detalle de la Visita ────────────────────────────────────
    sec += 1
    _titulo_seccion(doc, f"{sec}. Detalle de la Visita:")
    for linea in detalle_visita.split("\n"):
        doc.add_paragraph(linea)

    # ─── Sección 3: Metodología (opcional) ──────────────────────────────────
    if hay_inspeccion:
        sec += 1
        _titulo_seccion(doc, f"{sec}. Metodología y Hallazgos de la Inspección Técnica")

        doc.add_paragraph("Pruebas de verificación:").runs[0].bold = True

        for i, insp in enumerate(inspecciones, 1):
            subt = doc.add_paragraph()
            _keep_con_siguiente(subt)
            subt.paragraph_format.left_indent = Inches(0.25)
            subt.add_run(f"{i}. {insp['nombre']}").bold = True

            for linea in str(insp.get("texto") or "").split("\n"):
                p = doc.add_paragraph(linea)
                p.paragraph_format.left_indent = Inches(0.5)

            for dim in insp["dimensiones"]:
                p = doc.add_paragraph()
                _keep_con_siguiente(p)
                p.paragraph_format.left_indent = Inches(0.5)
                p.add_run(dim["titulo"]).underline = True
                for linea in str(dim.get("texto") or "").split("\n"):
                    pd_ = doc.add_paragraph(linea)
                    pd_.paragraph_format.left_indent = Inches(0.75)

    # ─── Sección Daños ──────────────────────────────────────────────────────
    sec += 1
    _titulo_seccion(doc, f"{sec}. Daños Identificados")

    if sin_danos:
        doc.add_paragraph("No se observan daños en el inmueble.")
    else:
        tabla_d = doc.add_table(rows=1, cols=4)
        tabla_d.style = "Table Grid"
        for i, h in enumerate(["Zona Afectada", "Ubicación", "Daño Observado", "Superficie Afectada"]):
            tabla_d.rows[0].cells[i].text = h

        for d in danos:
            row = tabla_d.add_row().cells
            row[0].text = d["zona"]
            row[1].text = d["ubicacion"]
            row[2].text = "\n".join(f"• {item['dano']}"       for item in d["items"])
            row[3].text = "\n".join(f"• {item['superficie']}" for item in d["items"])

    # ─── Sección Observaciones ───────────────────────────────────────────────
    sec += 1
    _titulo_seccion(doc, f"{sec}. Observaciones:")
    for obs in observaciones:
        doc.add_paragraph(f"• {obs}")

    # ─── Fotos (sin numeral, siempre al final) ───────────────────────────────
    p_fotos = _titulo_seccion(doc, "Fotos Riesgo Asegurado")
    p_fotos.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if imagen_error:
        doc.add_paragraph(imagen_error)
    elif not imagenes_data:
        doc.add_paragraph("No se encontraron imágenes válidas.")
    else:
        i, total = 0, len(imagenes_data)
        while i < total:
            if i + 1 < total:
                # Foto y pie de foto en la misma celda → nunca se separan
                tbl = doc.add_table(rows=1, cols=2)
                tbl.style = "Table Grid"
                tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                tbl.autofit = False
                for j in range(2):
                    cell = tbl.rows[0].cells[j]
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    # Foto
                    p_img = cell.paragraphs[0]
                    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    img_buf = imagenes_data[i + j]["bytes"]
                    img_buf.seek(0)
                    p_img.add_run().add_picture(img_buf, height=Inches(3.2))
                    # Pie de foto en la misma celda
                    desc = imagenes_data[i + j]["descripcion"].replace("_", " ").capitalize()
                    p_cap = cell.add_paragraph()
                    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_cap.add_run(f"Imagen {i + j + 1}: {desc}").italic = True
                i += 2
            else:
                tbl = doc.add_table(rows=1, cols=1)
                tbl.style = "Table Grid"
                tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                tbl.autofit = False
                cell = tbl.rows[0].cells[0]
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                # Foto
                p_img = cell.paragraphs[0]
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                img_buf = imagenes_data[i]["bytes"]
                img_buf.seek(0)
                p_img.add_run().add_picture(img_buf, height=Inches(4.0))
                # Pie de foto en la misma celda
                desc = imagenes_data[i]["descripcion"].replace("_", " ").capitalize()
                p_cap = cell.add_paragraph()
                p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_cap.add_run(f"Imagen {i + 1}: {desc}").italic = True
                i += 1
            doc.add_paragraph("")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
